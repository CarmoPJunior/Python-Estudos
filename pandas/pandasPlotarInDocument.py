import pandas as pd
import json


requisicao_itens = [
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"hostid": "10594",
				"host": "BARRETOS"
			}
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"itemid": "48316",
				"clock": "1654570816",
				"value": "33.903636",
				"ns": "318370897"
			}
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"itemid": "48328",
				"clock": "1654570828",
				"value": "94.34388274508902",
				"ns": "334136024"
			}
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"itemid": "233404",
				"clock": "1654620605",
				"value": "{\"Servername\":\"BARRETOS\\\\BARRETOS\",\"Hostname\":\"BARRETOS\",\"SQLVersion\":\"Microsoft SQL Server 2017 (RT\",\"OSVersion\":\"Windows Server 2012 R2 Datacenter 6.3 \\u003cX64\\u003e (Build 9600: ) (Hypervisor)\\n\",\"ServerIp\":\"190.1.1.10\",\"InstancePort\":62209,\"Instance\":\"BARRETOS\",\"ProductVersion\":\"14.0.3421.10\",\"Edition\":\"Standard Edition (64-bit)\",\"ProductLevel\":\"RTM\",\"ProductUpdateLevel\":\"CU27\",\"Collation\":\"SQL_Latin1_General_CP1_CI_AI\",\"ServerMemory\":65535,\"TargetServerMemory\":57344,\"MaxDop\":8,\"RemoteDAC\":1,\"CTFP\":30,\"IsVirtualMachine\":true,\"SystemModel\":\"System Manufacturer: \\u0027VMware, Inc.\\u0027, System Model: \\u0027VMware Virtual Platform\\u0027.\",\"CPUModel\":\"Intel(R) Xeon(R) Gold 5220R CPU @ 2.20GHz\",\"CPUCount\":24,\"NUMANodesCount\":4,\"LogicalCPUsCount\":24,\"OfflineCPUsCount\":0,\"SQLServerStartTime\":\"Apr 15 2022  1:17PM\",\"Uptime\":\"53d 0h 33m\",\"DatabasesCount\":38,\"ServiceName\":\"BARRETOS\",\"ServiceAccount\":\"MINERVA\\\\SVC_SQLSERVICES\",\"InstantFileInitialization\":true,\"SQLDiskCount\":2,\"SQLDiskDetails\":\"D:\\\\ com tamanho total de 2201GB e 823GB livres.,C:\\\\ com tamanho total de 159GB e 35GB livres.\",\"SQLTotalSizeGB\":1229.56,\"IsClustered\":0,\"IsAlwaysOnAGEnabled\":0,\"LogShippingDatabases\":0,\"MirroringDatabases\":0,\"ReplicationDatabases\":11,\"TraceFlags\":\"[1117,1118,2371,3226,9481]\"}",
				"ns": "883333417"
			}
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"itemid": "233405",
				"clock": "1654620607",
				"value": "{\"Servername\":\"BARRETOS\\\\PABX\",\"Hostname\":\"BARRETOS\",\"SQLVersion\":\"Microsoft SQL Server 2008 R2 (SP3)\",\"OSVersion\":\"Windows NT 6.3 \\u003cX64\\u003e (Build 9600: ) (Hypervisor)\\n\",\"ServerIp\":null,\"InstancePort\":null,\"Instance\":\"PABX\",\"ProductVersion\":\"10.50.6000.34\",\"Edition\":\"Standard Edition (64-bit)\",\"ProductLevel\":\"SP3\",\"ProductUpdateLevel\":null,\"Collation\":\"SQL_Latin1_General_CP1_CI_AI\",\"ServerMemory\":65535,\"TargetServerMemory\":512,\"MaxDop\":0,\"RemoteDAC\":0,\"CTFP\":5,\"IsVirtualMachine\":false,\"SystemModel\":\"System Manufacturer: \\u0027VMware, Inc.\\u0027, System Model: \\u0027VMware Virtual Platform\\u0027.\",\"CPUModel\":\"Intel(R) Xeon(R) Gold 5220R CPU @ 2.20GHz\",\"CPUCount\":24,\"NUMANodesCount\":2,\"LogicalCPUsCount\":24,\"OfflineCPUsCount\":0,\"SQLServerStartTime\":\"Apr 15 2022  1:17PM\",\"Uptime\":\"53d 0h 33m\",\"DatabasesCount\":6,\"ServiceName\":\"PABX\",\"ServiceAccount\":\"MINERVA\\\\SVC_SQLSERVICES\",\"InstantFileInitialization\":null,\"SQLDiskCount\":2,\"SQLDiskDetails\":null,\"SQLTotalSizeGB\":80.59,\"IsClustered\":0,\"IsAlwaysOnAGEnabled\":null,\"LogShippingDatabases\":0,\"MirroringDatabases\":0,\"ReplicationDatabases\":0,\"TraceFlags\":null}",
				"ns": "133247735"
			}
		],
		"id": 1
	}
]

dados = []

dados.append(requisicao_itens[3])
df = pd.json_normalize(dados)
df = df.explode('result')
df = df['result'].apply(pd.Series)
df = df['value']
df = df[0]
# df = df.strip()

# print(df)
# print(type(df))

df2 = pd.read_json(df, orient ='index')
print(df2)

print(100 *"-")
for index, row in df2.iterrows():
    print(index, row[0])

# for row in df2.itertuples():
#     print(row)
 



# a_json = json.loads(df)
# print(a_json)
# print(100*'-')
# print(type(a_json))

# for key, value in a_json.items():
	# print(key , ':' , value)

# Converter 
# print(100*'-')
# dataframe = pd.DataFrame.from_dict(a_json, orient="index")
# print(dataframe)


# pip install python-docx
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('./Pixel-House.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Item'
hdr_cells[1].text = 'Valor'


for index, row in df2.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(index)
    row_cells[1].text =  str(row[0])

document.add_page_break()

document.save('demo.docx')