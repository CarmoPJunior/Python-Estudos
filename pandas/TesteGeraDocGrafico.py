import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from PIL import Image
import io

requisicao_itens = [
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"hostid": "10594",
				"host": "BARRETOS"
			}
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
                "itemid": "48325",
                "clock": "1655245216",
                "value": "15",
                "ns": "740506090"
            },
            {
                "itemid": "48325",
                "clock": "1655245105",
                "value": "35",
                "ns": "245404386"
            }
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
                "itemid": "48316",
                "clock": "1655245216",
                "value": "67.096152",
                "ns": "354689327"
            },
            {
                "itemid": "48316",
                "clock": "1655245105",
                "value": "51.523859",
                "ns": "666694162"
            }
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"itemid": "233404",
				"clock": "1654620605",
				"value": "{\"Servername\":\"BARRETOS\\\\BARRETOS\",\"Hostname\":\"BARRETOS\",\"SQLVersion\":\"Microsoft SQL Server 2017 (RT\",\"OSVersion\":\"Windows Server 2012 R2 Datacenter 6.3 \\u003cX64\\u003e (Build 9600: ) (Hypervisor)\\n\",\"ServerIp\":\"190.1.1.10\",\"InstancePort\":62209,\"Instance\":\"BARRETOS\",\"ProductVersion\":\"14.0.3421.10\",\"Edition\":\"Standard Edition (64-bit)\",\"ProductLevel\":\"RTM\",\"ProductUpdateLevel\":\"CU27\",\"Collation\":\"SQL_Latin1_General_CP1_CI_AI\",\"ServerMemory\":65535,\"TargetServerMemory\":57344,\"MaxDop\":8,\"RemoteDAC\":1,\"CTFP\":30,\"IsVirtualMachine\":true,\"SystemModel\":\"System Manufacturer: \\u0027VMware, Inc.\\u0027, System Model: \\u0027VMware Virtual Platform\\u0027.\",\"CPUModel\":\"Intel(R) Xeon(R) Gold 5220R CPU @ 2.20GHz\",\"CPUCount\":24,\"NUMANodesCount\":4,\"LogicalCPUsCount\":24,\"OfflineCPUsCount\":0,\"SQLServerStartTime\":\"Apr 15 2022  1:17PM\",\"Uptime\":\"53d 0h 33m\",\"DatabasesCount\":38,\"ServiceName\":\"BARRETOS\",\"ServiceAccount\":\"MINERVA\\\\SVC_SQLSERVICES\",\"InstantFileInitialization\":true,\"SQLDiskCount\":2,\"SQLDiskDetails\":\"D:\\\\ com tamanho total de 2201GB e 823GB livres.,C:\\\\ com tamanho total de 159GB e 35GB livres.\",\"SQLTotalSizeGB\":1229.56,\"IsClustered\":0,\"IsAlwaysOnAGEnabled\":0,\"LogShippingDatabases\":0,\"MirroringDatabases\":0,\"ReplicationDatabases\":11,\"TraceFlags\":\"[1117,1118,2371,3226,9481]\"}",
				"ns": "883333417"
			}
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"itemid": "233405",
				"clock": "1654620607",
				"value": "{\"Servername\":\"BARRETOS\\\\PABX\",\"Hostname\":\"BARRETOS\",\"SQLVersion\":\"Microsoft SQL Server 2008 R2 (SP3)\",\"OSVersion\":\"Windows NT 6.3 \\u003cX64\\u003e (Build 9600: ) (Hypervisor)\\n\",\"ServerIp\":null,\"InstancePort\":null,\"Instance\":\"PABX\",\"ProductVersion\":\"10.50.6000.34\",\"Edition\":\"Standard Edition (64-bit)\",\"ProductLevel\":\"SP3\",\"ProductUpdateLevel\":null,\"Collation\":\"SQL_Latin1_General_CP1_CI_AI\",\"ServerMemory\":65535,\"TargetServerMemory\":512,\"MaxDop\":0,\"RemoteDAC\":0,\"CTFP\":5,\"IsVirtualMachine\":false,\"SystemModel\":\"System Manufacturer: \\u0027VMware, Inc.\\u0027, System Model: \\u0027VMware Virtual Platform\\u0027.\",\"CPUModel\":\"Intel(R) Xeon(R) Gold 5220R CPU @ 2.20GHz\",\"CPUCount\":24,\"NUMANodesCount\":2,\"LogicalCPUsCount\":24,\"OfflineCPUsCount\":0,\"SQLServerStartTime\":\"Apr 15 2022  1:17PM\",\"Uptime\":\"53d 0h 33m\",\"DatabasesCount\":6,\"ServiceName\":\"PABX\",\"ServiceAccount\":\"MINERVA\\\\SVC_SQLSERVICES\",\"InstantFileInitialization\":null,\"SQLDiskCount\":2,\"SQLDiskDetails\":null,\"SQLTotalSizeGB\":80.59,\"IsClustered\":0,\"IsAlwaysOnAGEnabled\":null,\"LogShippingDatabases\":0,\"MirroringDatabases\":0,\"ReplicationDatabases\":0,\"TraceFlags\":null}",
				"ns": "133247735"
			}
		],
		"id": 1
	}
]

# Trata as informações
infoHost =  pd.json_normalize(requisicao_itens[0]['result'])
infoHost = infoHost['host'][0]

infoHostCPU = pd.json_normalize(requisicao_itens[1]['result'])
infoHostCPU['value']  =infoHostCPU['value'] .astype('float64')
infoHostCPU["clock"]  = pd.to_datetime(infoHostCPU["clock"] , unit='s', origin='unix') 
# infoHostCPU['clock'] = infoHostCPU['clock'].dt.strftime('%d-%m-%Y %H:%M')

infoHostMemory = pd.json_normalize(requisicao_itens[2]['result'])
infoHostMemory['value']  =infoHostMemory['value'] .astype('float64')
infoHostMemory["clock"]  = pd.to_datetime(infoHostMemory["clock"] , unit='s', origin='unix') 
# infoHostMemory['clock'] = infoHostMemory['clock'].dt.strftime('%d-%m-%Y %H:%M')


df = pd.json_normalize(requisicao_itens[3])
df = df.explode('result')
df = df['result'].apply(pd.Series)
df = df['value']
df = df[0]
df2 = pd.read_json(df, orient ='index')
# print(df2)


# Plota o Gráfico
plt.ylabel('%')
plt.xlabel('Memória X Processador')
plt.title(infoHost)
plt.yticks([0,10,20,30, 40, 50, 60, 70, 80, 90, 100])
plt.legend()
ax = plt.gca()

infoHostMemory.plot(kind='line',x='clock',y='value', color='blue', label="Memória", ax=ax)
infoHostCPU.plot(kind='line',x='clock',y='value', color='red', label="Processador", ax=ax)
# fig = plt.gcf() #Aramazena o gráfico para ser impresso depois
# plt.show()    


# figureBuffer = io.BytesIO()
# plt.savefig(figureBuffer, format='png')
# figureBuffer.seek(0)
# im = Image.open(figureBuffer)
# im.show()
# figureBuffer.close()


# Armazena o gráfico em bytes
figure = io.BytesIO(); 
plt.savefig(figure); 
figure.seek(0); 


document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture(figure, width=Inches(3.25))

table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Item'
hdr_cells[1].text = 'Valor'


for index, row in df2.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(index)
    row_cells[1].text =  str(row[0])

document.add_page_break()

document.save('demo.docx')







