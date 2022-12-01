import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from PIL import Image
import io

requisicao_itens = [
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"itemid": "239262",
				"clock": "1654570800",
				"num": "60",
				"value_min": "5.572707",
				"value_avg": "8.193133433333328",
				"value_max": "20.008807"
			},
			{
				"itemid": "239262",
				"clock": "1654574400",
				"num": "60",
				"value_min": "5.324038",
				"value_avg": "8.484855916666666",
				"value_max": "24.728953"
			},
			{
				"itemid": "239262",
				"clock": "1654578000",
				"num": "60",
				"value_min": "5.955536",
				"value_avg": "15.341690600000002",
				"value_max": "31.92217"
			},
			{
				"itemid": "239262",
				"clock": "1654581600",
				"num": "60",
				"value_min": "14.169615",
				"value_avg": "27.954414283333325",
				"value_max": "42.027634"
			}
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"itemid": "239273",
				"clock": "1654570800",
				"num": "60",
				"value_min": "94.56781675720187",
				"value_avg": "94.84520423740675",
				"value_max": "95.36534505249273"
			},
			{
				"itemid": "239273",
				"clock": "1654574400",
				"num": "60",
				"value_min": "95.18632110405353",
				"value_avg": "95.28401368096226",
				"value_max": "95.38797711282979"
			},
			{
				"itemid": "239273",
				"clock": "1654578000",
				"num": "60",
				"value_min": "94.2424384212156",
				"value_avg": "94.7461520679347",
				"value_max": "95.29078502016291"
			},
			{
				"itemid": "239273",
				"clock": "1654581600",
				"num": "60",
				"value_min": "95.07169451665769",
				"value_avg": "95.43211449823771",
				"value_max": "96.1070233595407"
			}
		],
		"id": 1
	},
	{
		"jsonrpc": "2.0",
		"result": [
			{
				"itemid": "233404",
				"clock": "1654620605",
				"value": "{\"Servername\":\"BARRETOS\\\\BARRETOS\",\"Hostname\":\"BARRETOS\",\"SQLVersion\":\"Microsoft SQL Server 2017 (RT\",\"OSVersion\":\"Windows Server 2012 R2 Datacenter 6.3 \\u003cX64\\u003e (Build 9600: ) (Hypervisor)\\n\",\"ServerIp\":\"190.1.1.10\",\"InstancePort\":62209,\"Instance\":\"BARRETOS\",\"ProductVersion\":\"14.0.3421.10\",\"Edition\":\"Standard Edition (64-bit)\",\"ProductLevel\":\"RTM\",\"ProductUpdateLevel\":\"CU27\",\"Collation\":\"SQL_Latin1_General_CP1_CI_AI\",\"ServerMemory\":65535,\"TargetServerMemory\":57344,\"MaxDop\":8,\"RemoteDAC\":1,\"CTFP\":30,\"IsVirtualMachine\":true,\"SystemModel\":\"System Manufacturer: \\u0027VMware, Inc.\\u0027, System Model: \\u0027VMware Virtual Platform\\u0027.\",\"CPUModel\":\"Intel(R) Xeon(R) Gold 5220R CPU @ 2.20GHz\",\"CPUCount\":24,\"NUMANodesCount\":4,\"LogicalCPUsCount\":24,\"OfflineCPUsCount\":0,\"SQLServerStartTime\":\"Apr 15 2022  1:17PM\",\"Uptime\":\"53d 0h 33m\",\"DatabasesCount\":38,\"ServiceName\":\"BARRETOS\",\"ServiceAccount\":\"MINERVA\\\\SVC_SQLSERVICES\",\"InstantFileInitialization\":true,\"SQLDiskCount\":2,\"SQLDiskDetails\":\"D:\\\\ com tamanho total de 2201GB e 823GB livres.,C:\\\\ com tamanho total de 159GB e 35GB livres.\",\"SQLTotalSizeGB\":1229.56,\"IsClustered\":0,\"IsAlwaysOnAGEnabled\":0,\"LogShippingDatabases\":0,\"MirroringDatabases\":0,\"ReplicationDatabases\":11,\"TraceFlags\":\"[1117,1118,2371,3226,9481]\"}",
				"ns": "883333417"
			},
			{
				"itemid": "233404",
				"clock": "1654707005",
				"value": "{\"Servername\":\"BARRETOS\\\\BARRETOS\",\"Hostname\":\"BARRETOS\",\"SQLVersion\":\"Microsoft SQL Server 2017 (RT\",\"OSVersion\":\"Windows Server 2012 R2 Datacenter 6.3 \\u003cX64\\u003e (Build 9600: ) (Hypervisor)\\n\",\"ServerIp\":\"190.1.1.10\",\"InstancePort\":62209,\"Instance\":\"BARRETOS\",\"ProductVersion\":\"14.0.3421.10\",\"Edition\":\"Standard Edition (64-bit)\",\"ProductLevel\":\"RTM\",\"ProductUpdateLevel\":\"CU27\",\"Collation\":\"SQL_Latin1_General_CP1_CI_AI\",\"ServerMemory\":65535,\"TargetServerMemory\":57344,\"MaxDop\":8,\"RemoteDAC\":1,\"CTFP\":30,\"IsVirtualMachine\":true,\"SystemModel\":\"System Manufacturer: \\u0027VMware, Inc.\\u0027, System Model: \\u0027VMware Virtual Platform\\u0027.\",\"CPUModel\":\"Intel(R) Xeon(R) Gold 5220R CPU @ 2.20GHz\",\"CPUCount\":24,\"NUMANodesCount\":4,\"LogicalCPUsCount\":24,\"OfflineCPUsCount\":0,\"SQLServerStartTime\":\"Apr 15 2022  1:17PM\",\"Uptime\":\"54d 0h 33m\",\"DatabasesCount\":38,\"ServiceName\":\"BARRETOS\",\"ServiceAccount\":\"MINERVA\\\\SVC_SQLSERVICES\",\"InstantFileInitialization\":true,\"SQLDiskCount\":2,\"SQLDiskDetails\":\"D:\\\\ com tamanho total de 2201GB e 822GB livres.,C:\\\\ com tamanho total de 159GB e 35GB livres.\",\"SQLTotalSizeGB\":1230.07,\"IsClustered\":0,\"IsAlwaysOnAGEnabled\":0,\"LogShippingDatabases\":0,\"MirroringDatabases\":0,\"ReplicationDatabases\":11,\"TraceFlags\":\"[1117,1118,2371,3226,9481]\"}",
				"ns": "706881593"
			},
			{
				"itemid": "233404",
				"clock": "1654793405",
				"value": "{\"Servername\":\"BARRETOS\\\\BARRETOS\",\"Hostname\":\"BARRETOS\",\"SQLVersion\":\"Microsoft SQL Server 2017 (RT\",\"OSVersion\":\"Windows Server 2012 R2 Datacenter 6.3 \\u003cX64\\u003e (Build 9600: ) (Hypervisor)\\n\",\"ServerIp\":\"190.1.1.10\",\"InstancePort\":62209,\"Instance\":\"BARRETOS\",\"ProductVersion\":\"14.0.3421.10\",\"Edition\":\"Standard Edition (64-bit)\",\"ProductLevel\":\"RTM\",\"ProductUpdateLevel\":\"CU27\",\"Collation\":\"SQL_Latin1_General_CP1_CI_AI\",\"ServerMemory\":65535,\"TargetServerMemory\":57344,\"MaxDop\":8,\"RemoteDAC\":1,\"CTFP\":30,\"IsVirtualMachine\":true,\"SystemModel\":\"System Manufacturer: \\u0027VMware, Inc.\\u0027, System Model: \\u0027VMware Virtual Platform\\u0027.\",\"CPUModel\":\"Intel(R) Xeon(R) Gold 5220R CPU @ 2.20GHz\",\"CPUCount\":24,\"NUMANodesCount\":4,\"LogicalCPUsCount\":24,\"OfflineCPUsCount\":0,\"SQLServerStartTime\":\"Apr 15 2022  1:17PM\",\"Uptime\":\"55d 0h 33m\",\"DatabasesCount\":38,\"ServiceName\":\"BARRETOS\",\"ServiceAccount\":\"MINERVA\\\\SVC_SQLSERVICES\",\"InstantFileInitialization\":true,\"SQLDiskCount\":2,\"SQLDiskDetails\":\"C:\\\\ com tamanho total de 159GB e 35GB livres.,D:\\\\ com tamanho total de 2201GB e 821GB livres.\",\"SQLTotalSizeGB\":1231.08,\"IsClustered\":0,\"IsAlwaysOnAGEnabled\":0,\"LogShippingDatabases\":0,\"MirroringDatabases\":0,\"ReplicationDatabases\":11,\"TraceFlags\":\"[1117,1118,2371,3226,9481]\"}",
				"ns": "724085044"
			},
			{
				"itemid": "233404",
				"clock": "1654879805",
				"value": "{\"Servername\":\"BARRETOS\\\\BARRETOS\",\"Hostname\":\"BARRETOS\",\"SQLVersion\":\"Microsoft SQL Server 2017 (RT\",\"OSVersion\":\"Windows Server 2012 R2 Datacenter 6.3 \\u003cX64\\u003e (Build 9600: ) (Hypervisor)\\n\",\"ServerIp\":\"190.1.1.10\",\"InstancePort\":62209,\"Instance\":\"BARRETOS\",\"ProductVersion\":\"14.0.3421.10\",\"Edition\":\"Standard Edition (64-bit)\",\"ProductLevel\":\"RTM\",\"ProductUpdateLevel\":\"CU27\",\"Collation\":\"SQL_Latin1_General_CP1_CI_AI\",\"ServerMemory\":65535,\"TargetServerMemory\":57344,\"MaxDop\":8,\"RemoteDAC\":1,\"CTFP\":30,\"IsVirtualMachine\":true,\"SystemModel\":\"System Manufacturer: \\u0027VMware, Inc.\\u0027, System Model: \\u0027VMware Virtual Platform\\u0027.\",\"CPUModel\":\"Intel(R) Xeon(R) Gold 5220R CPU @ 2.20GHz\",\"CPUCount\":24,\"NUMANodesCount\":4,\"LogicalCPUsCount\":24,\"OfflineCPUsCount\":0,\"SQLServerStartTime\":\"Apr 15 2022  1:17PM\",\"Uptime\":\"56d 0h 33m\",\"DatabasesCount\":38,\"ServiceName\":\"BARRETOS\",\"ServiceAccount\":\"MINERVA\\\\SVC_SQLSERVICES\",\"InstantFileInitialization\":true,\"SQLDiskCount\":2,\"SQLDiskDetails\":\"C:\\\\ com tamanho total de 159GB e 35GB livres.,D:\\\\ com tamanho total de 2201GB e 821GB livres.\",\"SQLTotalSizeGB\":1231.59,\"IsClustered\":0,\"IsAlwaysOnAGEnabled\":0,\"LogShippingDatabases\":0,\"MirroringDatabases\":0,\"ReplicationDatabases\":11,\"TraceFlags\":\"[1117,1118,2371,3226,9481]\"}",
				"ns": "766097168"
			}
		],
		"id": 1
	}
]

# Trata as informações
# infoHost =  pd.json_normalize(requisicao_itens[0]['result'])
# infoHost = infoHost['host'][0]

infoHostCPU = pd.json_normalize(requisicao_itens[1]['result'])
infoHostCPU['value_avg']  =infoHostCPU['value_avg'] .astype('float64')
infoHostCPU["clock"]  = pd.to_datetime(infoHostCPU["clock"] , unit='s', origin='unix') 
# infoHostCPU['clock'] = infoHostCPU['clock'].dt.strftime('%d-%m-%Y %H:%M')

infoHostMemory = pd.json_normalize(requisicao_itens[0]['result'])
infoHostMemory['value_avg']  =infoHostMemory['value_avg'] .astype('float64')
infoHostMemory["clock"]  = pd.to_datetime(infoHostMemory["clock"] , unit='s', origin='unix') 
# infoHostMemory['clock'] = infoHostMemory['clock'].dt.strftime('%d-%m-%Y %H:%M')


df = pd.json_normalize(requisicao_itens[2])
df = df.explode('result')
df = df['result'].apply(pd.Series)
df = df['value']
df = df[0]
print(type (df))


json = []

for i in df:
	json.append(pd.read_json(i, orient ='index'))
# 	print("row: ", row)
# df2 = pd.read_json(df, orient ='index')
# print(df2)


# Plota o Gráfico
plt.ylabel('%')
plt.xlabel('Memória X Processador')
plt.title("infoHost")
plt.yticks([0,10,20,30, 40, 50, 60, 70, 80, 90, 100])
plt.legend()
ax = plt.gca()

infoHostMemory.plot(kind='line',x='clock',y='value_avg', color='blue', label="Memória", ax=ax)
infoHostCPU.plot(kind='line',x='clock',y='value_avg', color='red', label="Processador", ax=ax)
# fig = plt.gcf() #Aramazena o gráfico para ser impresso depois
# plt.show()    


# # figureBuffer = io.BytesIO()
# # plt.savefig(figureBuffer, format='png')
# # figureBuffer.seek(0)
# # im = Image.open(figureBuffer)
# # im.show()
# # figureBuffer.close()


# # Armazena o gráfico em bytes
figure = io.BytesIO(); 
plt.savefig(figure); 
figure.seek(0); 


document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture(figure, width=Inches(3.25))

table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Item'
hdr_cells[1].text = 'Valor'


# for index, row in df2.iterrows():
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(index)
#     row_cells[1].text =  str(row[0])

for dfTemp in json:
	for index, row in dfTemp.iterrows():
		row_cells = table.add_row().cells
		row_cells[0].text = str(index)
		row_cells[1].text =  str(row[0])

document.add_page_break()

document.save('demo.docx')







